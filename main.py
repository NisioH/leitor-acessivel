import flet as ft
import os
import sys
import tempfile
from io import BytesIO
import pytesseract
from PIL import Image
from gtts import gTTS
import PyPDF2
import docx
import pandas as pd
import requests
import base64
import shutil

# Pasta global para uploads
UPLOAD_DIR = os.path.abspath("uploads")

def check_tesseract_available():
    tesseract_bin = shutil.which("tesseract")
    if tesseract_bin:
        try:
            pytesseract.pytesseract.tesseract_cmd = tesseract_bin
            pytesseract.get_tesseract_version()
            return True
        except:
            return False
    return False


def ocr_online(image_bytes, language='por'):
    
    try:
        base64_image = base64.b64encode(image_bytes).decode()

        url = "https://api.ocr.space/parse/image"

        payload = {
            'base64Image': f'data:image/png;base64,{base64_image}',
            'language': language,
            'isOverlayRequired': False,
            'detectOrientation': True,
            'scale': True,
            'OCREngine': 2,  # Engine 2 é melhor para português
        }

        # Chave API gratuita (pública, limitada)
        headers = {
            'apikey': 'K87899142388957',  # Chave demo gratuita
        }

        response = requests.post(url, data=payload, headers=headers, timeout=30)
        result = response.json()

        if result.get('IsErroredOnProcessing'):
            raise Exception(f"Erro no OCR: {result.get('ErrorMessage', 'Erro desconhecido')}")

        parsed_results = result.get('ParsedResults', [])
        if parsed_results:
            text = parsed_results[0].get('ParsedText', '')
            return text
        else:
            raise Exception("Nenhum texto encontrado na imagem")

    except Exception as e:
        raise Exception(f"Erro no OCR online: {str(e)}")



def main(page: ft.Page):
    page.title = "Leitor Acessível"
    page.theme_mode = ft.ThemeMode.LIGHT
    page.padding = 20
    page.scroll = "adaptive"

    if page.web:
        page.theme = ft.Theme(
            page_transitions={"android": "zoom", "ios": "cupertino", "macos": "none"}
        )

    audio_player = None

    text_field = ft.TextField(
        label="Texto para leitura",
        multiline=True,
        min_lines=5,
        max_lines=20,
        hint_text="O texto extraído aparecerá aqui...",
        expand=False,
        text_size=14,
    )

    # Indicador de carregamento
    loading_indicator = ft.ProgressBar(
        visible=False,
        color=ft.Colors.BLUE_400,
        bgcolor=ft.Colors.BLUE_50
    )

    audio_added = False

    def on_file_result(e: ft.FilePickerResultEvent):
        if not e.files:
            return

        file_info = e.files[0]

        if page.web and not file_info.path and not getattr(file_info, "bytes", None):
            loading_indicator.visible = True
            page.update()

            upload_url = page.get_upload_url(file_info.name, 600)
            file_picker.upload([
                ft.FilePickerUploadFile(
                    file_info.name,
                    upload_url=upload_url,
                )
            ])
            return

        process_file(file_info)

    def on_upload_progress(e: ft.FilePickerUploadEvent):
        if e.progress == 1.0:
            class FakeFileInfo:
                def __init__(self, name, path):
                    self.name = name
                    self.path = path
                    self.bytes = None
            
            uploaded_file_path = os.path.join(UPLOAD_DIR, e.file_name)
            process_file(FakeFileInfo(e.file_name, uploaded_file_path))
        elif e.error:
            loading_indicator.visible = False
            text_field.value = f"Erro no upload: {e.error}"
            page.update()

    def process_file(file_info):
        loading_indicator.visible = True
        text_field.value = ""  # Limpa o texto anterior
        page.update()

        file_bytes = getattr(file_info, "bytes", None)
        file_name = getattr(file_info, "name", "uploaded_file")
        file_path = file_info.path

        try:
            extracted_text = ""
            if file_path:
                ext = os.path.splitext(file_path)[1].lower()
            else:
                ext = os.path.splitext(file_name)[1].lower()

            if ext == ".txt":
                if file_bytes is not None:
                    try:
                        extracted_text = file_bytes.decode("utf-8")
                    except Exception:
                        extracted_text = file_bytes.decode("latin-1", errors="replace")
                elif file_path:
                    with open(file_path, "r", encoding="utf-8") as f:
                        extracted_text = f.read()
                else:
                    raise ValueError("O conteúdo do arquivo não pôde ser lido (caminho e bytes ausentes).")

            elif ext == ".pdf":
                if file_bytes is not None:
                    pdf_reader = PyPDF2.PdfReader(BytesIO(file_bytes))
                elif file_path:
                    pdf_reader = PyPDF2.PdfReader(file_path)
                else:
                    raise ValueError("O PDF não pôde ser carregado.")
                
                for page_pdf in pdf_reader.pages:
                    extracted_text += (page_pdf.extract_text() or "") + "\n"

            elif ext == ".docx":
                if file_bytes is not None:
                    doc = docx.Document(BytesIO(file_bytes))
                elif file_path:
                    doc = docx.Document(file_path)
                else:
                    raise ValueError("O DOCX não pôde ser carregado.")
                
                for para in doc.paragraphs:
                    extracted_text += para.text + "\n"

            elif ext in [".xlsx", ".xls"]:
                if file_bytes is not None:
                    df = pd.read_excel(BytesIO(file_bytes))
                elif file_path:
                    df = pd.read_excel(file_path)
                else:
                    raise ValueError("O arquivo Excel não pôde ser carregado.")
                
                extracted_text = df.to_string(index=False)

            elif ext in [".png", ".jpg", ".jpeg", ".bmp"]:
                if file_bytes is not None:
                    img = Image.open(BytesIO(file_bytes))
                    img_bytes_for_ocr = file_bytes
                elif file_path:
                    img = Image.open(file_path)
                    # Converter para bytes para uso no OCR online
                    img_buffer = BytesIO()
                    img.save(img_buffer, format='PNG')
                    img_bytes_for_ocr = img_buffer.getvalue()
                else:
                    raise ValueError("A imagem não pôde ser carregada.")

                extracted_text = ""
                ocr_method_used = ""
                
                if check_tesseract_available():
                    try:
                        text_field.value = "🔍 Extraindo texto (OCR local)..."
                        page.update()
                        extracted_text = pytesseract.image_to_string(img, lang='por')
                    except Exception:
                        extracted_text = ""

                if not extracted_text.strip():
                    try:
                        text_field.value = "🌐 Extraindo texto (OCR online)..."
                        page.update()
                        extracted_text = ocr_online(img_bytes_for_ocr, language='por')
                        ocr_method_used = " ✓ [OCR Online]"
                    except Exception as ocr_ex:
                        raise RuntimeError(f"Erro no OCR: {str(ocr_ex)}")
                else:
                    ocr_method_used = " ✓ [OCR Local]"

                if extracted_text.strip():
                    extracted_text = f"{extracted_text}\n\n---\n{ocr_method_used}"

            # Atualizar o campo de texto com o resultado (seja imagem ou documento)
            text_field.value = extracted_text if extracted_text.strip() else "Nenhum texto encontrado."
                
        except Exception as ex:
            import traceback
            error_details = traceback.format_exc()
            print(f"Erro ao processar arquivo: {ex}\n{error_details}")
            text_field.value = f"Erro ao processar arquivo: {ex}"

        loading_indicator.visible = False
        page.update()

    file_picker = ft.FilePicker(on_result=on_file_result, on_upload=on_upload_progress)
    page.overlay.append(file_picker)

    def open_camera(e):

        file_picker.pick_files(
            allow_multiple=False,
            file_type=ft.FilePickerFileType.IMAGE,
        )

    def download_audio(e):
        if audio_player is not None and audio_player.src:
            src_path = audio_player.src
            
            # Se for web, tenta baixar via URL
            if page.web:
                download_url = f"uploads/{src_path}" if not src_path.startswith("/") else src_path
                page.launch_url(download_url)
            else:

                page.launch_url(src_path)

    download_button = ft.ElevatedButton(
        "📥 Baixar MP3",
        icon=ft.Icons.DOWNLOAD,
        visible=False,
        on_click=download_audio,
        tooltip="Baixar ou compartilhar áudio",
        style=ft.ButtonStyle(
            bgcolor=ft.Colors.BLUE_700,
            color=ft.Colors.WHITE,
        )
    )

    audio_status = ft.Text(
        "",
        size=12,
        color=ft.Colors.GREEN_700,
        weight="bold",
        text_align=ft.TextAlign.CENTER,
        visible=False
    )

    def convert_and_play(e):
        nonlocal audio_added, audio_player
        text_value = (text_field.value or "").strip()
        if not text_value:
            text_field.value = "Nenhum texto para converter."
            page.update()
            return

        loading_indicator.visible = True
        download_button.visible = False
        audio_status.visible = False
        page.update()

        try:
            import time
            timestamp = int(time.time())
            audio_name = f"leitura_{timestamp}.mp3"

            # Pasta de uploads (funciona no PC e será incluída no APK)
            if not os.path.exists(UPLOAD_DIR):
                try:
                    os.makedirs(UPLOAD_DIR)
                except Exception:
                    pass
            
            audio_path_save = os.path.join(UPLOAD_DIR, audio_name)

            # Limpar áudios antigos
            for f in os.listdir(UPLOAD_DIR):
                if f.startswith("leitura_") and f.endswith(".mp3"):
                    try:
                        os.remove(os.path.join(UPLOAD_DIR, f))
                    except Exception:
                        pass

            tts = gTTS(text=text_value, lang='pt')
            tts.save(audio_path_save)

            # O segredo é usar o prefixo '/' para arquivos em assets_dir/upload_dir
            audio_url = f"/{audio_name}"

            if audio_player is None:
                audio_player = ft.Audio(src=audio_url, autoplay=False)
                page.overlay.append(audio_player)
                audio_added = True
            else:
                audio_player.src = audio_url
                if not audio_added:
                    page.overlay.append(audio_player)
                    audio_added = True
            
            page.update() # Atualiza para garantir que o player carregou o novo src

            audio_player.volume = 1

            download_button.visible = True
            audio_status.value = "🔊 Áudio pronto! Tocando..."
            audio_status.visible = True
            page.update()

            success = False
            for _ in range(5):
                try:
                    audio_player.play()
                    success = True
                    break
                except:
                    time.sleep(0.2)

            if not success:
                audio_status.value = "⚠️ Áudio gerado, mas não tocou automaticamente. Use o botão de download."
                audio_status.color = ft.Colors.ORANGE_700
        except Exception as ex:
            audio_status.value = f"❌ Erro: {ex}"
            audio_status.color = ft.Colors.RED_700
            audio_status.visible = True
        
        loading_indicator.visible = False
        page.update()

    tesseract_available = check_tesseract_available()
    ocr_info = ft.Container(
        content=ft.Row([
            ft.Icon(ft.Icons.INFO_OUTLINE, color=ft.Colors.BLUE_700, size=20),
            ft.Text(
                f"{'✅ OCR Local + Online' if tesseract_available else '🌐 OCR Online'} disponível",
                size=12,
                color=ft.Colors.BLUE_700,
                weight="bold"
            )
        ], tight=True),
        bgcolor=ft.Colors.BLUE_50,
        border_radius=8,
        padding=10,
    )

    page.add(
        ft.Column([
            ft.Text("Leitor Acessível", size=28, weight="bold", text_align=ft.TextAlign.CENTER),
            ft.Text(
                "Transforme fotos e arquivos em voz",
                size=14,
                text_align=ft.TextAlign.CENTER,
                color=ft.Colors.GREY_700
            ),
            ft.Divider(),
            ocr_info,
            ft.Row([
                ft.Container(
                    content=ft.ElevatedButton(
                        "📷 Foto",
                        icon=ft.Icons.CAMERA_ALT,
                        on_click=open_camera,
                        height=50,
                        tooltip="Tire uma foto ou escolha da galeria",
                        style=ft.ButtonStyle(
                            bgcolor=ft.Colors.BLUE_700,
                            color=ft.Colors.WHITE,
                        )
                    ),
                    expand=1,
                ),
                ft.Container(
                    content=ft.ElevatedButton(
                        "📁 Arquivo",
                        icon=ft.Icons.UPLOAD_FILE,
                        on_click=lambda _: file_picker.pick_files(
                            allowed_extensions=["txt", "png", "jpg", "jpeg", "pdf", "docx", "xlsx", "xls", "bmp"]
                        ),
                        height=50,
                        style=ft.ButtonStyle(
                            bgcolor=ft.Colors.PURPLE_700,
                            color=ft.Colors.WHITE,
                        )
                    ),
                    expand=1,
                ),
            ], spacing=10),
            ft.Column([
                ft.Text(
                    "Imagens • PDF • DOCX • XLSX • TXT",
                    size=12,
                    color=ft.Colors.GREY_600,
                    text_align=ft.TextAlign.CENTER
                ),
                ft.Text(
                    "💡 Dica: No botão Foto, escolha 'Câmera' para tirar na hora",
                    size=11,
                    color=ft.Colors.BLUE_600,
                    text_align=ft.TextAlign.CENTER,
                    italic=True,
                    visible=page.web
                ),
            ], spacing=5, horizontal_alignment=ft.CrossAxisAlignment.CENTER),
            loading_indicator,
            ft.Container(
                content=text_field,
                expand=True,
            ),
            ft.Container(
                content=ft.Column([
                    audio_status,
                    ft.ElevatedButton(
                        "🔊 Ouvir Texto",
                        icon=ft.Icons.PLAY_ARROW,
                        bgcolor=ft.Colors.GREEN_700,
                        color=ft.Colors.WHITE,
                        on_click=convert_and_play,
                        width=300,
                        height=50,
                    ),
                    download_button
                ], horizontal_alignment=ft.CrossAxisAlignment.CENTER, spacing=10),
                alignment=ft.alignment.center
            )
        ],
        horizontal_alignment=ft.CrossAxisAlignment.STRETCH,
        spacing=15,
        scroll=ft.ScrollMode.AUTO
        )
    )

def get_local_ip():
    import socket
    import subprocess

    try:
        result = subprocess.run(['hostname', '-I'], capture_output=True, text=True)
        ips = result.stdout.strip().split()
        for ip in ips:
            if ip.startswith('192.168.') or ip.startswith('10.'):
                return ip
        if ips:
            return ips[0]
    except:
        pass

    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        ip = s.getsockname()[0]
        s.close()
        if not ip.startswith('127.'):
            return ip
    except:
        pass

    return None


if __name__ == "__main__":
    # Garantir caminho absoluto para evitar erros de diretório
    if not os.path.exists(UPLOAD_DIR):
        try:
            os.makedirs(UPLOAD_DIR)
        except:
            pass

    os.environ["FLET_SECRET_KEY"] = "leitor_acessivel_secret_2024"

    print("=" * 60)
    print("🚀 Iniciando Leitor Acessível...")
    print("=" * 60)
    print(f"📁 Pasta de uploads: {UPLOAD_DIR}")

    local_ip = get_local_ip()

    if local_ip:
        print(f"\n📱 ACESSE NO CELULAR (App Flet):")
        print(f"   http://{local_ip}:8550")
        print(f"\n💻 Ou no navegador do computador:")
        print(f"   http://localhost:8550")
    else:
        print(f"\n⚠️  Não foi possível detectar o IP automaticamente.")

    print("\n" + "=" * 60 + "\n")

    try:
        ft.app(
            target=main,
            view=ft.AppView.WEB_BROWSER,
            upload_dir=UPLOAD_DIR,
            assets_dir=UPLOAD_DIR,
            port=8550,
        )
    except Exception as e:
        print(f"⚠️ Erro ao iniciar: {e}")
        ft.app(target=main, upload_dir=UPLOAD_DIR, assets_dir=UPLOAD_DIR)
